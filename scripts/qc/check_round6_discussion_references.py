#!/usr/bin/env python3
from pathlib import Path
import re, sys
import pandas as pd
from docx import Document

ROOT = Path(__file__).resolve().parents[2]
DOC = ROOT / "outputs/jtm_manuscript_round6_discussion_enhanced.docx"
checks = []
def add(name, ok):
    checks.append((name, bool(ok)))
def text_doc(path):
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)
txt = text_doc(DOC) if DOC.exists() else ""
disc = txt.split("Discussion",1)[1].split("Conclusions",1)[0] if "Discussion" in txt and "Conclusions" in txt else ""
refs = txt.split("References",1)[1].split("Figures",1)[0] if "References" in txt and "Figures" in txt else ""
ref_nums = sorted(set(int(x) for x in re.findall(r"^(\d+)\.", refs, re.M)))
cite_set = set()
for block in re.findall(r"\[([0-9,\-\s]+)\]", txt):
    for part in block.replace(" ", "").split(","):
        if not part:
            continue
        if "-" in part:
            a, b = part.split("-", 1)
            cite_set.update(range(int(a), int(b) + 1))
        else:
            cite_set.add(int(part))
cite_nums = sorted(cite_set)
add("Discussion word count reasonable", 1200 <= len(disc.split()) <= 1800)
add("principal findings included", "principal finding" in disc.lower() and "430" in disc and "17" in disc)
add("vascular/ECM axis discussed", "extracellular-matrix" in disc.lower() and "VWF" in disc and "COL5A1" in disc)
add("myeloid/complement axis discussed", "myeloid" in disc.lower() and "complement" in disc.lower())
add("scRNA/bulk/protein differences explained", all(x in disc for x in ["scRNA", "Bulk", "Protein"]))
add("protein FDR negative explicit", "no candidate met protein FDR<0.05" in txt or "no candidate met FDR<0.05" in txt)
add("previous integration studies compared", "public-data integration" in disc.lower() and "prediction model" in disc.lower())
add("genetics/TWAS explained", "nearest-gene" in disc.lower() and "colocalization" in disc.lower() and "per-SNP" in disc)
add("drug context separated from treatment", "therapeutic context" in disc and "treatment-ready" not in disc.lower())
add("validation roadmap present", "validation roadmap" in disc.lower() and "spatial transcriptomics" in disc.lower())
add("430 upstream unrecoverable acknowledged", "not fully recoverable" in disc.lower() or "not fully reconstructable" in disc.lower())
add("17 rule retrospective acknowledged", "retrospective" in disc.lower() and "preregistered" in disc.lower())
add("no independent holdout acknowledged", "no fully independent molecular hold-out" in disc.lower())
add("pathway concordance not gene validation", "pathway-level concordance" in disc and "do not constitute independent gene-level validation" in disc)
bad = ["proved", "confirms", "key drivers", "clinically actionable", "repurposing candidate", "promising treatment"]
add("unsupported causal/clinical wording absent", not any(b in disc.lower() for b in bad))
add("literature search log exists", (ROOT/"reports/jtm_round6_literature_search_log.xlsx").exists())
add("DOI field audited", (ROOT/"reports/reference_add_remove_log.tsv").exists() and pd.read_csv(ROOT/"reports/reference_add_remove_log.tsv", sep="\t")["doi_verified"].isin(["yes","not_applicable"]).all())
add("all references cited", set(ref_nums).issubset(set(cite_nums)))
add("no missing references for citations", set(cite_nums).issubset(set(ref_nums)))
add("method audit exists", (ROOT/"reports/method_reference_audit_round6.tsv").exists())
if (ROOT/"reports/method_reference_audit_round6.tsv").exists():
    m = pd.read_csv(ROOT/"reports/method_reference_audit_round6.tsv", sep="\t")
    add("upstream-only methods moved", m[m["used_in_round5"].eq("no")]["keep_remove_or_move"].isin(["move","remove"]).all())
else:
    add("upstream-only methods moved", False)
add("no duplicate reference numbers", len(ref_nums) == len(set(ref_nums)))
add("citation numbering continuous", ref_nums == list(range(1, max(ref_nums)+1)) if ref_nums else False)
add("reference format includes DOI for most", refs.lower().count("doi:") >= 28)
add("Discussion not Methods repeat", disc.lower().count("dirichlet") <= 1 and disc.lower().count("checksum") <= 1)
add("Conclusion consistent", "independent molecular validation" in txt and "treatment-ready" in txt)
add("independent validation not claimed", "independent validation was achieved" not in txt.lower())
add("drug not called treatment candidate", "treatment candidate" not in txt.lower())
add("nearest-gene not causal", "nearest-gene assignment does not identify a causal gene" in disc.lower())
add("protein nominal not validation", "protein validation" not in txt.lower())
failed = [x for x in checks if not x[1]]
report = ROOT / "reports/jtm_round6_discussion_reference_qc.md"
report.write_text("# Round6 Discussion/reference QC\n\n" + "\n".join(f"- {'PASS' if ok else 'FAIL'}: {name}" for name, ok in checks) + "\n", encoding="utf-8")
block = ROOT / "reports/jtm_round6_submission_blockers.md"
block.write_text("# Round6 submission blockers\n\n" + ("- Technical QC failures: %d.\n" % len(failed) if failed else "- Technical QC: no failures in implemented checks.\n") + "- Author metadata, funding, conflicts, repository URL/DOI and final access dates remain author-side blockers.\n", encoding="utf-8")
print(f"Round6 QC: {len(checks)-len(failed)} PASS, {len(failed)} FAIL")
for name, _ in failed:
    print("FAIL:", name)
sys.exit(1 if failed else 0)
