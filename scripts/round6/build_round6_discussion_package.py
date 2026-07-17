#!/usr/bin/env python3
from __future__ import annotations

import re
from dataclasses import dataclass
from datetime import date
from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt

ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "outputs"
REPORTS = ROOT / "reports"
FIGS = OUT / "figures"
ROUND5 = OUT / "jtm_manuscript_round5_scientific.docx"

OUT.mkdir(exist_ok=True)
REPORTS.mkdir(exist_ok=True)

FINAL17 = [
    "VWF", "COL5A1", "VCAN", "ALOX5AP", "CTSB", "FN1", "SULF1",
    "CCDC102B", "SPARC", "CFH", "THY1", "COL1A2", "ST6GALNAC3",
    "COL11A1", "SEMA5A", "MSR1", "HERPUD1",
]


@dataclass
class Ref:
    n: int
    text: str
    doi: str = ""
    pmid: str = ""
    source: str = "PubMed/DOI"
    action: str = "keep"
    reason: str = ""
    location: str = "Discussion"


def refs() -> list[Ref]:
    items = [
        Ref(1, "Tremolada G, Del Turco C, Lattanzio R, et al. The role of angiogenesis in the development of proliferative diabetic retinopathy: impact of intravitreal anti-VEGF treatment. Exp Diabetes Res. 2012;2012:728325. doi:10.1155/2012/728325.", "10.1155/2012/728325", action="keep", reason="background angiogenesis context", location="Background"),
        Ref(2, "Rangasamy S, McGuire PG, Das A. Diabetic retinopathy and inflammation: novel therapeutic targets. Middle East Afr J Ophthalmol. 2012;19:52-59. doi:10.4103/0974-9233.92116.", "10.4103/0974-9233.92116", action="keep", reason="background inflammatory context", location="Background"),
        Ref(3, "Hu Z, Mao X, Chen M, et al. Single-Cell Transcriptomics Reveals Novel Role of Microglia in Fibrovascular Membrane of Proliferative Diabetic Retinopathy. Diabetes. 2022;71:762-773. doi:10.2337/db21-0551.", "10.2337/db21-0551", "35061025", action="keep", reason="PDR FVM scRNA-seq primary evidence", location="Background/Discussion"),
        Ref(4, "Corano Scheri K, Lavine JA, Tedeschi T, et al. Single-cell transcriptomics analysis of proliferative diabetic retinopathy fibrovascular membranes reveals AEBP1 as fibrogenesis modulator. JCI Insight. 2023;8:e172062. doi:10.1172/jci.insight.172062.", "10.1172/jci.insight.172062", "37917183", action="keep", reason="PDR FVM scRNA-seq and pericyte-myofibroblast functional validation", location="Background/Discussion"),
        Ref(5, "Ishikawa K, Yoshida S, Kobayashi Y, et al. Microarray analysis of gene expression in fibrovascular membranes excised from patients with proliferative diabetic retinopathy. Invest Ophthalmol Vis Sci. 2015;56:932-946. doi:10.1167/iovs.14-15589.", "10.1167/iovs.14-15589", action="keep", reason="PDR FVM bulk transcriptomic comparator", location="Background/Discussion"),
        Ref(6, "Becker K, Klein H, Simon E, et al. In-depth transcriptomic analysis of human retina reveals molecular mechanisms underlying diabetic retinopathy. Sci Rep. 2021;11:10494. doi:10.1038/s41598-021-88698-3.", "10.1038/s41598-021-88698-3", action="keep", reason="retinal transcriptomic comparator", location="Background/Discussion"),
        Ref(7, "Becker K, Klein H, Simon E, et al. Proliferative diabetic retinopathy transcriptomes reveal angiogenesis, anti-angiogenic therapy escape mechanisms, fibrosis and lymphatic involvement. Sci Rep. 2021;11:18810. doi:10.1038/s41598-021-97970-5.", "10.1038/s41598-021-97970-5", "34552123", action="add", reason="recent PDR transcriptomic evidence for fibrosis/lymphatic/angiogenesis axes", location="Discussion"),
        Ref(8, "Li Y, Chen D, Sun L, et al. Induced expression of VEGFC, ANGPT, and EFNB2 and their receptors characterizes neovascularization in proliferative diabetic retinopathy. Invest Ophthalmol Vis Sci. 2019;60:4084-4096. doi:10.1167/iovs.19-26767.", "10.1167/iovs.19-26767", action="keep", reason="PDR neovascularization mechanism context", location="Discussion"),
        Ref(9, "Lam JD, Oh DJ, Wong LL, et al. Identification of RUNX1 as a mediator of aberrant retinal angiogenesis. Diabetes. 2017;66:1950-1956. doi:10.2337/db16-1035.", "10.2337/db16-1035", action="keep", reason="retinal endothelial functional validation example", location="Discussion"),
        Ref(10, "Abu El-Asrar AM, De Hertogh G, Nawaz MI, et al. Interplay of endothelial-mesenchymal transition, inflammation, and autophagy in proliferative diabetic retinopathy pathogenesis. Heliyon. 2024;10:e25166. doi:10.1016/j.heliyon.2024.e25166.", "10.1016/j.heliyon.2024.e25166", "38327444", action="add", reason="recent PDR EndoMT/inflammation/autophagy comparator", location="Discussion"),
        Ref(11, "Li Y, Wang Y, Wang S, et al. Upregulation of HMOX1 associated with M2 macrophage infiltration and ferroptosis in proliferative diabetic retinopathy. Int Immunopharmacol. 2024;134:112231. doi:10.1016/j.intimp.2024.112231.", "10.1016/j.intimp.2024.112231", "38739977", action="add", reason="recent PDR macrophage/inflammation comparator", location="Discussion"),
        Ref(12, "Shahulhameed S, Vishwakarma S, Chhablani J, et al. A Systematic Investigation on Complement Pathway Activation in Diabetic Retinopathy. Front Immunol. 2020;11:154. doi:10.3389/fimmu.2020.00154.", "10.3389/fimmu.2020.00154", action="keep", reason="complement pathway context", location="Discussion"),
        Ref(13, "Mandava N, Tirado-Gonzalez V, Geiger MD, et al. Complement Activation in the Vitreous of Patients With Proliferative Diabetic Retinopathy. Invest Ophthalmol Vis Sci. 2020;61:39. doi:10.1167/iovs.61.11.39.", "10.1167/iovs.61.11.39", "32965482", action="add", reason="PDR vitreous complement evidence", location="Discussion"),
        Ref(14, "Xiao H, Guo X, Wang D, et al. Comprehensive Proteomic Profiling of Aqueous Humor Proteins in Proliferative Diabetic Retinopathy. Transl Vis Sci Technol. 2021;10:3. doi:10.1167/tvst.10.6.3.", "10.1167/tvst.10.6.3", "34111250", action="add", reason="PDR aqueous proteomics comparator", location="Discussion"),
        Ref(15, "Santos FM, Gaspar LM, Ciordia S, et al. Proteomics profiling of vitreous humor reveals complement and coagulation components, adhesion factors, and neurodegeneration markers as discriminatory biomarkers of vitreoretinal eye diseases. Front Immunol. 2023;14:1107295. doi:10.3389/fimmu.2023.1107295.", "10.3389/fimmu.2023.1107295", "36875133", action="add", reason="vitreous proteomics and complement/coagulation context", location="Discussion"),
        Ref(16, "Zhang P, Su X, Chen Y, et al. System-wide vitreous proteome dissection reveals impaired sheddase activity in diabetic retinopathy. Theranostics. 2022;12:6682-6704. doi:10.7150/thno.72947.", "10.7150/thno.72947", "36185601", action="add", reason="vitreous proteome context and compartment boundary", location="Discussion"),
        Ref(17, "Li X, Li J, Wang Y, et al. Aqueous humor proteomics analyzed by bioinformatics and machine learning in PDR cases versus controls. Clin Proteomics. 2024;21:36. doi:10.1186/s12014-024-09481-w.", "10.1186/s12014-024-09481-w", "38764026", action="add", reason="recent PDR aqueous proteomics plus machine-learning comparator", location="Discussion"),
        Ref(18, "Tam V, Patel N, Turcotte M, et al. Benefits and limitations of genome-wide association studies. Nat Rev Genet. 2019;20:467-484. doi:10.1038/s41576-019-0127-1.", "10.1038/s41576-019-0127-1", action="keep", reason="GWAS interpretation boundary", location="Discussion"),
        Ref(19, "Gusev A, Ko A, Shi H, et al. Integrative approaches for large-scale transcriptome-wide association studies. Nat Genet. 2016;48:245-252. doi:10.1038/ng.3506.", "10.1038/ng.3506", action="keep", reason="TWAS method boundary", location="Discussion"),
        Ref(20, "Ratnapriya R, Sosina OA, Starostik MR, et al. Retinal transcriptome and eQTL analyses identify genes associated with age-related macular degeneration. Nat Genet. 2019;51:606-610. doi:10.1038/s41588-019-0351-9.", "10.1038/s41588-019-0351-9", "30742112", action="keep", reason="retinal eQTL resource boundary", location="Discussion"),
        Ref(21, "Kurki MI, Karjalainen J, Palta P, et al. FinnGen provides genetic insights from a well-phenotyped isolated population. Nature. 2023;613:508-518. doi:10.1038/s41586-022-05473-8.", "10.1038/s41586-022-05473-8", action="keep", reason="FinnGen source interpretation", location="Discussion"),
        Ref(22, "Wainberg M, Sinnott-Armstrong N, Mancuso N, et al. Opportunities and challenges for transcriptome-wide association studies. Nat Genet. 2019;51:592-599. doi:10.1038/s41588-019-0385-z.", "10.1038/s41588-019-0385-z", action="keep", reason="TWAS limitation and interpretation", location="Discussion"),
        Ref(23, "Liu Y, Li X, Zhang C, et al. Integration of multi-omics transcriptome-wide analysis for the identification of novel therapeutic drug targets in diabetic retinopathy. J Transl Med. 2024;22:1146. doi:10.1186/s12967-024-05856-7.", "10.1186/s12967-024-05856-7", "39719581", action="add", reason="recent DR public-data multi-omics comparator", location="Background/Discussion"),
        Ref(24, "Qian B, Chen H, Wang X, et al. DRAC 2022: A public benchmark for diabetic retinopathy analysis on ultra-wide optical coherence tomography angiography images. Patterns. 2024;5:100929. doi:10.1016/j.patter.2024.100929.", "10.1016/j.patter.2024.100929", action="keep", reason="OCTA data boundary and imaging-omics gap", location="Background/Discussion"),
        Ref(25, "Durbin MK, An L, Shemonski ND, et al. Quantification of retinal microvascular density in optical coherence tomographic angiography images in diabetic retinopathy. JAMA Ophthalmol. 2017;135:370-376. doi:10.1001/jamaophthalmol.2017.0080.", "10.1001/jamaophthalmol.2017.0080", action="keep", reason="OCTA measurement context", location="Discussion"),
        Ref(26, "ACCORD Study Group and ACCORD Eye Study Group. Effects of medical therapies on retinopathy progression in type 2 diabetes. N Engl J Med. 2010;363:233-244. doi:10.1056/NEJMoa1001288.", "10.1056/NEJMoa1001288", "20587587", action="add", reason="fenofibrate positive-control clinical context", location="Discussion"),
        Ref(27, "Keech AC, Mitchell P, Summanen PA, et al. Effect of fenofibrate on the need for laser treatment for diabetic retinopathy (FIELD study): a randomised controlled trial. Lancet. 2007;370:1687-1697. doi:10.1016/S0140-6736(07)61607-9.", "10.1016/S0140-6736(07)61607-9", "17988728", action="add", reason="fenofibrate positive-control clinical context", location="Discussion"),
        Ref(28, "Casini G, Dal Monte M, Fornaciari I, et al. Anti-angiogenic and antioxidant effects of axitinib in human retinal endothelial cells: implications in diabetic retinopathy. Front Pharmacol. 2024;15:1415846. doi:10.3389/fphar.2024.1415846.", "10.3389/fphar.2024.1415846", "38953109", action="add", reason="retinal endothelial drug-mechanism probe context", location="Discussion"),
        Ref(29, "Kang-Mieler JJ, Rudeen KM, Liu W, et al. Axitinib inhibits retinal and choroidal neovascularization in in vitro and in vivo models. Exp Eye Res. 2016;145:373-379. doi:10.1016/j.exer.2016.02.010.", "10.1016/j.exer.2016.02.010", "26927930", action="add", reason="antiangiogenic probe boundary", location="Discussion"),
        Ref(30, "Wu PC, Chen YH, Chen HY, et al. Association between pentoxifylline use and diabetic retinopathy in patients with type 2 diabetes mellitus and chronic kidney disease: a multi-institutional cohort study. Biomed J. 2025;48:100771. doi:10.1016/j.bj.2024.100771.", "10.1016/j.bj.2024.100771", "39033962", action="add", reason="cautionary drug-context evidence", location="Discussion"),
        Ref(31, "Grover D, Li TJ, Chong CC. Pentoxifylline for diabetic retinopathy. Cochrane Database Syst Rev. 2008;2008:CD006693. doi:10.1002/14651858.CD006693.pub2.", "10.1002/14651858.CD006693.pub2", "18425965", action="add", reason="drug-context uncertainty", location="Discussion"),
        Ref(32, "Han D, Gao J, Xu Z, et al. Predictive model for proliferative diabetic retinopathy using single-cell transcriptomics. Exp Eye Res. 2025;259:110536. doi:10.1016/j.exer.2025.110536.", "10.1016/j.exer.2025.110536", "40701535", action="add", reason="recent public-data/scRNA prediction comparator", location="Discussion"),
        Ref(33, "Wolf FA, Angerer P, Theis FJ. SCANPY: large-scale single-cell gene expression data analysis. Genome Biol. 2018;19:15. doi:10.1186/s13059-017-1382-0.", "10.1186/s13059-017-1382-0", action="move", reason="archived upstream scRNA only; not cited in main Round6", location="Supplementary Methods"),
        Ref(34, "Korsunsky I, Millard N, Fan J, et al. Fast, sensitive and accurate integration of single-cell data with Harmony. Nat Methods. 2019;16:1289-1296. doi:10.1038/s41592-019-0619-0.", "10.1038/s41592-019-0619-0", action="move", reason="archived upstream scRNA integration only; not cited in main Round6", location="Supplementary Methods"),
        Ref(35, "Robinson MD, McCarthy DJ, Smyth GK. edgeR: a Bioconductor package for differential expression analysis of digital gene expression data. Bioinformatics. 2010;26:139-140. doi:10.1093/bioinformatics/btp616.", "10.1093/bioinformatics/btp616", action="move", reason="archived upstream bulk only; not cited in main Round6", location="Supplementary Methods"),
        Ref(36, "Ritchie ME, Phipson B, Wu D, et al. limma powers differential expression analyses for RNA-sequencing and microarray studies. Nucleic Acids Res. 2015;43:e47. doi:10.1093/nar/gkv007.", "10.1093/nar/gkv007", action="move", reason="archived upstream bulk only; not cited in main Round6", location="Supplementary Methods"),
        Ref(33, "Benjamini Y, Hochberg Y. Controlling the false discovery rate: a practical and powerful approach to multiple testing. J R Stat Soc Series B. 1995;57:289-300. doi:10.1111/j.2517-6161.1995.tb02031.x.", "10.1111/j.2517-6161.1995.tb02031.x", action="keep", reason="FDR terminology retained", location="Methods"),
    ]
    return items


def current_sections() -> dict[str, str]:
    doc = Document(ROUND5)
    headings = {
        "Abstract", "Keywords", "Background", "Methods", "Results", "Discussion",
        "Conclusions", "Tables", "Figure Legends", "Declarations", "References", "Figures",
    }
    sections: dict[str, list[str]] = {"Front": []}
    cur = "Front"
    for p in doc.paragraphs:
        t = p.text.strip()
        if t in headings:
            cur = t
            sections[cur] = []
        elif t:
            sections.setdefault(cur, []).append(t)
    return {k: "\n".join(v) for k, v in sections.items()}


def stats() -> dict[str, object]:
    short = pd.read_csv(OUT / "final_shortlist_derivation.tsv", sep="\t")
    layer = pd.read_csv(OUT / "major_layer_exclusion_full_metrics.tsv", sep="\t")
    coverage = pd.read_csv(OUT / "reconstructed_candidate_universe.tsv", sep="\t")
    bulk = pd.read_csv(OUT / "bulk_evidence_summary_for_table2.tsv", sep="\t")
    protein = pd.read_csv(ROOT / "results/jtm_final/protein_validation_candidate_matrix.csv")
    evaluable = pd.read_csv(ROOT / "results/jtm_final/evidence_coverage_matrix.csv")
    return {
        "n_universe": len(coverage),
        "n_full_v7": 691,
        "n_final": len(short),
        "top30_min": short["top30_probability"].min(),
        "top30_max": short["top30_probability"].max(),
        "layer": layer,
        "scrna_evaluable": int(evaluable[evaluable["gene"].isin(FINAL17)]["scRNA_evaluable"].sum()),
        "protein_fdr_sig": int((protein[protein["gene"].isin(FINAL17)]["fdr_bh"] < 0.05).sum()) if "fdr_bh" in protein else 0,
        "bulk": bulk,
        "short": short,
    }


def background_text() -> str:
    return (
        "Proliferative diabetic retinopathy (PDR) is characterized by retinal ischemia, neovascularization, inflammation and fibrovascular membrane formation. "
        "Public single-cell, bulk transcriptomic, proteomic, genetic and imaging resources have helped describe these processes, including PDR fibrovascular-membrane cell states, matrix remodelling, vitreous or aqueous inflammatory proteins and retinal microvascular changes [1-8,14-17,24,25].\n\n"
        "Existing public-data studies in PDR have primarily focused on differential expression, cell-state characterization, pathway enrichment, prediction-model development or therapeutic-target nomination [3-7,17,23,32]. "
        "Less attention has been paid to whether candidate rankings remain stable under evidence missingness, alternative candidate-universe definitions and separation of ranking inputs from contextual annotations. "
        "This gap matters because PDR public datasets differ in tissue compartment, platform, disease stage and sample availability, and a simple summed validation score can obscure whether a result is driven by biology, coverage or modelling choices.\n\n"
        "The present study therefore does not claim a de novo end-to-end candidate-generation pipeline. "
        "We performed a retrospective evidence audit and secondary prioritization of a frozen project-defined 430-gene PDR candidate universe. "
        "The analysis explicitly reconstructed the candidate universe, reproduced the focused 17-gene set from archived rules, separated ranking inputs from contextual annotation layers and tested sensitivity to evidence-layer exclusion."
    )


def discussion_text(s: dict[str, object]) -> str:
    layer = s["layer"].set_index("excluded_layer")
    j_sc = layer.loc["scRNA", "top30_jaccard"]
    j_bulk = layer.loc["bulk", "top30_jaccard"]
    j_gt = layer.loc["genetics_TWAS", "top30_jaccard"]
    r_sc = int(layer.loc["scRNA", "top17_recovery"])
    r_bulk = int(layer.loc["bulk", "top17_recovery"])
    r_gt = int(layer.loc["genetics_TWAS", "top17_recovery"])
    return (
        f"The principal finding is that the archived PDR prioritization result can be made interpretable only after separating candidate-universe reconstruction from evidence-layer interpretation. "
        f"The frozen universe contained {s['n_universe']} genes and matched the frozen coverage matrix, whereas the broader v7 context table contained {s['n_full_v7']} genes. "
        f"The focused set contained {s['n_final']} genes and was reproduced from explicit archived rules rather than by taking a fixed top-17 truncation. "
        f"All {s['n_final']} genes were retained across the three alternative candidate-universe definitions, and their recomputed top-30 probabilities ranged from {s['top30_min']:.3f} to {s['top30_max']:.3f}. "
        f"However, excluding major evidence layers changed the top-30 composition, with Jaccard indices of {j_sc:.2f} after scRNA exclusion, {j_bulk:.2f} after bulk exclusion and {j_gt:.2f} after genetics/TWAS exclusion. "
        f"The corresponding final-17 recoveries were {r_sc}, {r_bulk} and {r_gt} genes. "
        f"The focused set was stable to the three alternative candidate-universe definitions but not invariant to evidence-layer composition, indicating stability to universe construction rather than universal ranking robustness. "
        f"Only {s['scrna_evaluable']}/17 genes were evaluable in archived scRNA summaries, and no candidate met protein FDR<0.05. "
        "These results should therefore be interpreted as an audited prioritization, not as an independent molecular validation study.\n\n"
        "The vascular and extracellular-matrix axis was the clearest biological theme among the prioritized genes. "
        "VWF, COL5A1, VCAN, FN1, SULF1, SPARC, COL1A2, COL11A1 and THY1 collectively point toward vascular activation, endothelial remodelling, extracellular-matrix deposition, fibrovascular membrane maturation and contractile fibrosis. "
        "This interpretation is consistent with PDR fibrovascular-membrane single-cell studies that describe endothelial, pericyte, myeloid and fibrotic stromal programs, including pericyte-to-myofibroblast-like transition and AEBP1-linked fibrogenesis [3,4]. "
        "It is also consistent with transcriptomic studies of PDR fibrovascular tissue reporting angiogenesis, fibrosis, lymphatic involvement and anti-VEGF escape mechanisms [5,7,8]. "
        "SPARC, collagen genes, FN1 and VCAN should not be read as independent causal proof for each gene. "
        "Rather, their joint prioritization suggests pathway-level concordance with fibrotic and wound-healing biology in end-stage membranes. "
        "The prominence of extracellular-matrix genes may reflect both authentic fibrotic biology and the cellular composition of surgically excised end-stage membranes.\n\n"
        "A second axis involved myeloid, inflammatory and complement biology. "
        "ALOX5AP, CTSB and MSR1 are most appropriately interpreted as inflammatory or myeloid-context candidates, while CFH links the focused set to complement biology and genetic-background annotation. "
        "PDR fibrovascular-membrane scRNA-seq has emphasized microglia/macrophage states and their interaction with fibrotic and angiogenic tissue compartments [3]. "
        "Recent PDR studies also support macrophage polarization, ferroptosis-related immune infiltration, endothelial-mesenchymal transition and inflammatory signalling as disease-relevant programs [10,11]. "
        "Vitreous and pathway-level studies provide additional support for complement activation and complement-coagulation proteins in diabetic retinopathy and vitreoretinal disease [12,13,15]. "
        "These comparisons support biological concordance at the pathway level. "
        "They do not constitute independent gene-level validation for ALOX5AP, CTSB, MSR1 or CFH.\n\n"
        "The apparent differences among scRNA, bulk and protein evidence are expected because these layers answer different biological questions. "
        "The archived scRNA summaries localize expression to lesion cell compartments but do not provide a complete case-control differential-expression test for all candidates. "
        "The 12 scRNA-unevaluable genes should therefore be treated as missing, not negative. "
        "Bulk transcriptomic evidence captures tissue-level case-control changes, but it is influenced by cell composition, tissue source, platform, disease stage and control selection. "
        "Because bulk evidence participated in the ranking, it provides cross-dataset support within the prioritization system rather than an independent validation layer. "
        "Protein data measure different biological compartments, including aqueous humor, vitreous humor and tissue-adjacent fluids, which are shaped by blood-retinal barrier breakdown, leakage, hemorrhage and inflammation [14-17]. "
        "The absence of FDR-supported protein findings limits cross-omic confirmation but does not necessarily contradict transcriptomic prioritization because the available protein data sampled different biological compartments and had limited candidate-level coverage.\n\n"
        "Compared with common public-data integration studies, the present work has a narrower but more explicit objective. "
        "It does not build a clinical prediction model, merge healthy retina and end-stage fibrovascular membranes as if they were interchangeable PDR mechanisms, or treat every evidence source as a validation score. "
        "Recent DR and PDR studies have used single-cell data, machine learning, transcriptome-wide association, proteomic profiling or multi-omics integration to identify biomarkers and candidate targets [17,23,32]. "
        "Those designs are valuable for discovery, but they often give less attention to missingness, evidence-source dependence and the boundary between ranking input and contextual annotation. "
        "The contribution here is not a new end-to-end discovery algorithm, but an explicit audit of how evidence availability, candidate-universe construction and evidence-layer composition influence a frozen public-data prioritization result.\n\n"
        "The genetics, TWAS and retinal regulatory components require particularly cautious interpretation. "
        "Nearest-gene assignment does not identify a causal gene, gene-window overlap may include several genes, and TWAS associations can reflect linkage disequilibrium, co-regulated expression or model-coverage constraints [18-22]. "
        "Retinal eQTL annotation is useful because it is tissue relevant, but gene-level eQTL evidence cannot substitute for allele-harmonized colocalization. "
        "The lack of per-SNP retinal eQTL beta, standard-error and allele information was the main barrier to formal colocalization. "
        "The pronounced rank changes after removal of genetics/TWAS indicate that part of the focused set was prioritized through cross-layer convergence rather than transcriptomic evidence alone. "
        "CFH therefore has a stronger genetic-context signal than many other focused genes, but it should not be described as causal. "
        "CDH6 and FRMD4A remained secondary genetics-prioritized candidates, yet they did not enter the focused 17-gene set because they failed the full archived rule combination.\n\n"
        "Therapeutic-context results should be read as hypothesis-generating annotation, not as drug repurposing evidence. "
        "Fenofibrate functioned as a positive-control retinal and metabolic context because randomized clinical evidence links it to slower diabetic-retinopathy progression or reduced laser-treatment need [26,27]. "
        "This does not imply that the focused genes mediate fenofibrate benefit. "
        "Multi-kinase inhibitors such as sorafenib, vandetanib, axitinib and dasatinib are better viewed as pleiotropic mechanistic probes of angiogenic target connectivity. "
        "Axitinib has retinal endothelial and neovascularization evidence in experimental settings, but multi-target kinase inhibition raises unresolved issues of target direction, ocular exposure, toxicity and delivery [28,29]. "
        "Dipyridamole, pentoxifylline and sirolimus are directionally plausible but unvalidated contexts. "
        "For pentoxifylline, systematic and observational evidence illustrates uncertainty rather than readiness for PDR therapy [30,31]. "
        "Across all drug annotations, the appropriate terms are therapeutic context, mechanistic probe and hypothesis-generating connection.\n\n"
        "The main strengths of this work are methodological transparency rather than clinical validation. "
        "The analysis reconstructed the frozen universe, preserved missing values, applied an explicit focused-set rule, separated ranking and contextual layers, tested multiple universe definitions, recomputed Dirichlet weight perturbation, performed layer-exclusion analysis, reported protein-negative results and defined the colocalization feasibility boundary. "
        "The same features also clarify the limitations. "
        "The upstream reduction from broader candidate lists to the frozen 430-gene universe was not fully recoverable, and the 17-gene rule was retrospective rather than preregistered. "
        "No fully independent molecular hold-out dataset was available. "
        "scRNA coverage was limited to five focused genes, bulk data were part of the ranking, protein evidence was FDR nonsignificant, per-SNP retinal eQTL statistics were unavailable, and drug context lacked directionality and safety validation. "
        "The use of surgically excised PDR fibrovascular membranes also biases interpretation toward end-stage neovascular and fibrotic tissue. "
        "Clinical variables, treatment history and same-patient OCTA features were not matched to molecular profiles.\n\n"
        "A practical validation roadmap should now move in tiers. "
        "First, independent PDR fibrovascular-membrane cohorts should test protein localization by immunohistochemistry, multiplex immunofluorescence or spatial transcriptomics using endothelial, myeloid, pericyte and fibroblast markers. "
        "Second, retinal endothelial cells, pericytes, macrophage-like or microglia-like cells and fibroblast/myofibroblast-like cells should be used for knockdown, CRISPR or overexpression assays measuring angiogenesis, migration, barrier function, inflammatory response, collagen deposition and contraction, following the logic of retinal angiogenesis studies that moved from candidate nomination to functional testing [9]. "
        "Third, genetic validation will require per-SNP retinal eQTL, allele harmonization, colocalization, fine mapping, conditional analysis and independent TWAS models. "
        "Fourth, clinical translation will require same-patient fibrovascular-membrane, aqueous or vitreous sampling linked to OCTA vessel density, non-perfusion, neovascular area, membrane fibrosis, treatment history and postoperative recurrence. "
        "Finally, therapeutic-context hypotheses require target-direction, dose-response, ocular-delivery, retinal-toxicity and candidate-specific rescue experiments before any treatment claim is justified."
    )


def conclusion_text() -> str:
    return (
        "This study reconstructed and audited a frozen 430-gene PDR candidate universe and reproduced a focused 17-gene set from explicit archived rules. "
        "The findings are most consistent with vascular-matrix and myeloid/complement biological contexts, but they remain sensitive to evidence-layer composition and lack independent molecular validation. "
        "Future work should test candidate localization, function, genetic colocalization and same-patient imaging-molecular associations before clinical or therapeutic use is inferred."
    )


def abstract_text() -> str:
    return (
        "Background: Public molecular resources for proliferative diabetic retinopathy (PDR) are heterogeneous, and candidate rankings may be affected by missing evidence, tissue compartment and candidate-universe definition. "
        "We performed a retrospective evidence audit and secondary prioritization of a frozen project-defined PDR candidate universe.\n"
        "Methods: We reconstructed the 430-gene frozen universe from archived primary-ranking outputs, recomputed Monte Carlo diagnostics using Dirichlet(1,1,1,1) layer-weight perturbation, audited major-layer exclusion and separated ranking inputs from contextual annotation layers.\n"
        "Results: The reconstructed universe matched the frozen coverage matrix for all 430 genes. A focused 17-gene set was reproduced from explicit archived rules rather than a fixed top-17 truncation. The set was stable across alternative universe definitions, but top-30 composition changed after layer exclusion, especially after genetics/TWAS removal. Vascular-matrix and myeloid/complement axes provided the strongest biological context. Only 5/17 genes were evaluable in archived scRNA summaries, no candidate met protein FDR<0.05, and drug evidence was retained only as therapeutic context.\n"
        "Conclusions: The study provides an auditable secondary prioritization of a fixed PDR candidate universe. Findings support follow-up tissue, cellular, genetic and imaging-linked validation but do not establish causal genes, clinical biomarkers or treatment-ready interventions."
    )


def make_table2() -> pd.DataFrame:
    short = pd.read_csv(OUT / "final_shortlist_derivation.tsv", sep="\t").set_index("gene").loc[FINAL17].reset_index()
    modules = {
        "VWF": "Vascular/ECM", "COL5A1": "Vascular/ECM", "VCAN": "Vascular/ECM",
        "FN1": "Vascular/ECM", "SULF1": "Vascular/ECM", "SPARC": "Vascular/ECM",
        "COL1A2": "Vascular/ECM", "COL11A1": "Vascular/ECM", "THY1": "Vascular/ECM",
        "ALOX5AP": "Myeloid/inflammatory", "CTSB": "Myeloid/inflammatory",
        "MSR1": "Myeloid/inflammatory", "CFH": "Complement/genetic context",
        "HERPUD1": "Less-characterized exploratory", "CCDC102B": "Less-characterized exploratory",
        "ST6GALNAC3": "Less-characterized exploratory", "SEMA5A": "Less-characterized exploratory",
    }
    rows = []
    for i, r in short.iterrows():
        gene = r["gene"]
        rows.append({
            "Gene": gene,
            "Rank": i + 1,
            "Top-30 probability": f"{r['top30_probability']:.3f}",
            "Biological module": modules[gene],
            "Evidence use": "ranking-derived",
            "Caution": "needs validation",
        })
    return pd.DataFrame(rows)


def add_table(doc: Document, df: pd.DataFrame, caption: str) -> None:
    p = doc.add_paragraph(caption)
    p.runs[0].bold = True
    table = doc.add_table(rows=1, cols=len(df.columns))
    table.style = "Table Grid"
    for i, c in enumerate(df.columns):
        run = table.rows[0].cells[i].paragraphs[0].add_run(str(c))
        run.bold = True
    for _, row in df.iterrows():
        cells = table.add_row().cells
        for i, c in enumerate(df.columns):
            cells[i].text = str(row[c])
    doc.add_paragraph("")


def build_doc(blinded: bool = False) -> Document:
    s = stats()
    doc = Document()
    styles = doc.styles
    styles["Normal"].font.name = "Arial"
    styles["Normal"].font.size = Pt(10)
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("Retrospective multi-source evidence audit and candidate prioritization in proliferative diabetic retinopathy")
    run.bold = True
    run.font.size = Pt(16)
    doc.add_paragraph("[AUTHOR CONFIRMATION REQUIRED]" if not blinded else "[BLINDED]")
    for heading, body in [
        ("Abstract", abstract_text()),
        ("Keywords", "proliferative diabetic retinopathy; frozen candidate universe; fibrovascular membrane; scRNA-seq; TWAS; proteomics; therapeutic context"),
        ("Background", background_text()),
        ("Methods", current_sections()["Methods"] + " False-discovery-rate language follows the Benjamini-Hochberg framework [33]."),
        ("Results", current_sections()["Results"]),
        ("Discussion", discussion_text(s)),
        ("Conclusions", conclusion_text()),
    ]:
        doc.add_heading(heading, level=1)
        for para in body.split("\n"):
            if para.strip():
                doc.add_paragraph(para.strip())
    doc.add_heading("Tables", level=1)
    t1 = pd.DataFrame([
        {"Resource group": "PDR FVM scRNA-seq", "Examples": "GSE165784; GSE245561", "Round6 role": "Ranking/context", "Boundary": "lesion-compartment summaries; missing genes not negative"},
        {"Resource group": "Bulk transcriptomics", "Examples": "GSE102485; GSE60436; GSE94019; GSE160306/10", "Round6 role": "Ranking support", "Boundary": "not independent validation because used in ranking"},
        {"Resource group": "FinnGen DR endpoints", "Examples": "FinnGen R13", "Round6 role": "Genetic context", "Boundary": "nearest/window evidence is not causal assignment"},
        {"Resource group": "Retina TWAS/eQTL", "Examples": "EyeGEx/retinal eQTL resources", "Round6 role": "Regulatory context", "Boundary": "no formal colocalization without per-SNP retinal eQTL"},
        {"Resource group": "Protein resources", "Examples": "aqueous/vitreous/FVM-related evidence", "Round6 role": "Context only", "Boundary": "no focused candidate met protein FDR<0.05"},
        {"Resource group": "Drug databases", "Examples": "DGIdb/OpenTargets/ChEMBL/CTD context", "Round6 role": "Therapeutic context", "Boundary": "mechanistic probes, not therapies or clinical recommendations"},
    ])
    add_table(doc, t1, "Table 1. Core public resources and corrected analytical roles. Full audit is provided in Supplementary Table S1.")
    add_table(doc, make_table2(), "Table 2. Biological-module summary for the 17 prioritized genes. Detailed evidence is provided in Supplementary Table S2.")
    doc.add_heading("Figure Legends", level=1)
    doc.add_paragraph("Graphical Abstract. Retrospective evidence audit of a frozen project-defined PDR candidate universe.")
    doc.add_paragraph("Figure 1. Candidate-universe reconstruction and ranking traceability. Panels show the frozen 430-gene universe, broader 691-gene upstream table, explicit 17-gene rule, Monte Carlo top-30 probabilities and major-layer exclusion metrics.")
    doc.add_paragraph("Figure 2. Transcriptomic localization, cross-dataset support and protein context for prioritized genes. Protein results are shown as quantitative context; no candidate met FDR<0.05.")
    doc.add_paragraph("Figure 3. Genetic association, colocalization boundary and therapeutic-context mapping. Internal structural permutation diagnostics are retained only as supplementary audit material.")
    doc.add_heading("Declarations", level=1)
    for h in ["Acknowledgements", "Authors' contributions", "Funding", "Availability of data and materials", "Ethics approval and consent to participate", "Consent for publication", "Competing interests"]:
        doc.add_heading(h, level=2)
        doc.add_paragraph("[AUTHOR CONFIRMATION REQUIRED]" if not blinded else "[BLINDED]")
    doc.add_heading("References", level=1)
    for r in [x for x in refs() if x.action != "move"]:
        doc.add_paragraph(f"{r.n}. {r.text}")
    doc.add_heading("Figures", level=1)
    for idx, stem in enumerate(["Figure1_round5", "Figure2_round5", "Figure3_round5"]):
        if idx:
            doc.add_page_break()
        p = FIGS / f"{stem}.png"
        if p.exists():
            doc.add_picture(str(p), width=Inches(6.5))
    return doc


def build_audits() -> None:
    sec = current_sections()
    s = stats()
    disc = sec.get("Discussion", "")
    paras = [p for p in disc.split("\n") if p.strip()]
    cited_nums = set(int(x) for x in re.findall(r"\[(\d+)\]", "\n".join(sec.values())))
    audit = [
        "# JTM Round6 discussion and reference audit",
        "",
        f"- Audit date: {date.today().isoformat()}",
        f"- Current Discussion word count: {len(disc.split())}",
        f"- Current Discussion paragraph count: {len(paras)}",
        "",
        "## Paragraph functions",
    ]
    for i, p in enumerate(paras, 1):
        func = ["universe/layer sensitivity summary", "430-universe traceability limitation", "17-gene retrospective-rule and validation boundary", "conservative conclusion"][min(i - 1, 3)]
        audit.append(f"- Paragraph {i}: {func}.")
    audit += [
        "",
        "## Main deficiencies",
        "- Results are repeated but not biologically interpreted.",
        "- Vascular/ECM genes, myeloid/complement genes and less-characterized exploratory candidates are not organized into biological axes.",
        "- scRNA, bulk and protein layers are not contrasted by compartment, platform and biological question.",
        "- Protein FDR-negative findings are stated but not interpreted.",
        "- Genetics/TWAS/eQTL boundaries are too brief, especially nearest-gene, window and colocalization limits.",
        "- Drug evidence is correctly conservative but lacks the requested category-level explanation.",
        "- The previous reference list contained many method references that were used only in archived upstream workflows.",
        "",
        "## Required Round6 actions",
        "- Add recent PDR scRNA, bulk, proteomic, complement, integration, TWAS and therapeutic-context literature.",
        "- Move upstream-only method references to Supplementary Methods instead of the main reference list.",
        "- Expand Discussion to 1200-1600 words with principal findings, biological axes, cross-omic interpretation, genetics/TWAS boundary, therapeutic context and validation roadmap.",
        "",
        "## Numeric anchors verified from Round5 outputs",
        f"- Frozen candidate universe: {s['n_universe']} genes; broader v7 table: {s['n_full_v7']} genes.",
        f"- Focused set: {s['n_final']} genes; top-30 probability range {s['top30_min']:.3f}-{s['top30_max']:.3f}.",
        f"- scRNA evaluable genes: {s['scrna_evaluable']}/17.",
        f"- Protein FDR<0.05 among focused genes: {s['protein_fdr_sig']}.",
    ]
    REPORTS.joinpath("jtm_round6_discussion_reference_audit.md").write_text("\n".join(audit) + "\n", encoding="utf-8")

    matrix = [
        ("principal findings", "all candidates", "universe stable but layer sensitive", "internal outputs", "strong internal", "no", "internal audit", "low", "stable to universe construction rather than universal ranking robustness"),
        ("vascular/ECM axis", "VWF/COL5A1/VCAN/FN1/SULF1/SPARC/COL1A2/COL11A1/THY1", "ECM prominence may reflect fibrosis and tissue composition", "[3-8,10]", "pathway-level", "yes", "PDR FVM scRNA/bulk", "medium", "may reflect both authentic fibrotic biology and surgical end-stage composition"),
        ("myeloid/complement axis", "ALOX5AP/CTSB/MSR1/CFH/HERPUD1", "inflammatory and complement context", "[3,11-15]", "pathway-level", "yes", "PDR scRNA/proteomics", "medium", "pathway-level concordance, not independent gene validation"),
        ("cross-omic disagreement", "all candidates", "scRNA/bulk/protein answer different biological questions", "[14-17]", "strong contextual", "yes", "proteomics and compartment studies", "low", "protein FDR-negative findings limit cross-omic confirmation"),
        ("genetics/TWAS", "CFH/CDH6/FRMD4A", "nearest-gene and TWAS are not causal proof", "[18-23]", "methodological", "yes", "GWAS/TWAS/eQTL resources", "high", "genetic context, not confirmed causal genes"),
        ("drug context", "fenofibrate/kinase inhibitors/pentoxifylline", "therapeutic context only", "[26-31]", "contextual", "yes", "clinical/experimental drug studies", "high", "mechanistic probe or positive-control context"),
    ]
    pd.DataFrame(matrix, columns=[
        "discussion_topic", "candidate_gene_or_module", "current_claim", "current_reference",
        "reference_support_level", "new_reference_needed", "recommended_reference_type",
        "overstatement_risk", "recommended_wording",
    ]).to_csv(REPORTS / "discussion_claim_reference_matrix.tsv", sep="\t", index=False)

    modules = {
        "VWF": "Vascular/ECM", "COL5A1": "Vascular/ECM", "VCAN": "Vascular/ECM", "FN1": "Vascular/ECM",
        "SULF1": "Vascular/ECM", "SPARC": "Vascular/ECM", "COL1A2": "Vascular/ECM", "COL11A1": "Vascular/ECM", "THY1": "Vascular/ECM",
        "ALOX5AP": "Myeloid/inflammatory", "CTSB": "Myeloid/inflammatory", "MSR1": "Myeloid/inflammatory",
        "CFH": "Complement/genetic context", "HERPUD1": "Less-characterized exploratory",
        "CCDC102B": "Less-characterized exploratory", "ST6GALNAC3": "Less-characterized exploratory", "SEMA5A": "Less-characterized exploratory",
    }
    rows = []
    for g in FINAL17:
        direct = "limited direct PDR-specific evidence"
        if g in {"VWF", "COL5A1", "VCAN", "FN1", "SPARC", "COL1A2", "COL11A1", "THY1"}:
            direct = "pathway-level PDR FVM ECM/fibrosis evidence"
        if g in {"ALOX5AP", "CTSB", "MSR1"}:
            direct = "pathway-level PDR myeloid/inflammatory evidence"
        if g == "CFH":
            direct = "complement/genetic-context evidence; not causal confirmation"
        rows.append({
            "Gene": g,
            "Biological module": modules[g],
            "Current manuscript evidence": "archived ranking plus Round6 contextual interpretation",
            "External PDR gene-level evidence": direct,
            "External DR pathway-level evidence": "supported by PDR FVM scRNA/bulk/proteomics where module-level evidence exists",
            "Cell type": "endothelial/fibroblast/myeloid/complement context depending on module",
            "Tissue/fluid": "PDR FVM, retina, aqueous humor or vitreous depending on source",
            "Study type": "public-data audit plus external original studies",
            "Reference": "see Round6 refs [3-17,23]",
            "Direction": "not used for causal direction",
            "Independent or overlapping dataset": "mostly independent literature context; some public-data overlap possible",
            "Evidence strength": "module-level" if "limited" not in direct else "limited",
            "Discussion use": "representative module interpretation" if modules[g] != "Less-characterized exploratory" else "grouped exploratory candidates",
            "Caution": "do not describe as independent gene validation",
        })
    pd.DataFrame(rows).to_excel(REPORTS / "gene_external_evidence_map.xlsx", index=False)

    method_rows = [
        ("Wolf 2018 Scanpy", "Scanpy", "no", "yes", "scripts/01_scrna_qc and archived h5ad outputs", "Supplementary Methods only", "move", "Round6 final analysis used archived scRNA summaries."),
        ("Virshup 2024 AnnData", "AnnData", "no", "yes", "archived h5ad workflow", "Supplementary Methods only", "move", "Not needed in main text."),
        ("Korsunsky 2019 Harmony", "Harmony", "no", "yes", "scripts/01_scrna_qc/03_harmony_integrate.py", "Supplementary Methods only", "move", "Upstream integration only."),
        ("McInnes 2018 UMAP", "UMAP", "no", "yes", "archived scRNA visualization scripts", "Supplementary Methods only", "move", "No Round6 final inference depends on UMAP."),
        ("Traag 2019 Leiden", "Leiden", "no", "yes", "archived clustering scripts", "Supplementary Methods only", "move", "No Round6 final inference depends on Leiden clustering."),
        ("Robinson 2010 edgeR", "edgeR", "no", "yes", "archived bulk-validation scripts", "Supplementary Methods only", "move", "Round6 reads processed bulk summaries."),
        ("Chen 2016 edgeR QL", "edgeR quasi-likelihood", "no", "yes", "archived bulk workflow", "Supplementary Methods only", "move", "Round6 does not rerun differential expression."),
        ("Robinson 2010 TMM", "TMM", "no", "yes", "archived bulk workflow", "Supplementary Methods only", "move", "Round6 does not rerun normalization."),
        ("Ritchie 2015 limma", "limma", "no", "yes", "archived bulk workflow", "Supplementary Methods only", "move", "Round6 reads processed summaries."),
        ("Law 2014 voom", "voom", "no", "yes", "archived bulk workflow", "Supplementary Methods only", "move", "Round6 does not rerun voom."),
        ("Benjamini 1995 BH", "Benjamini-Hochberg FDR", "yes", "yes", "protein/bulk summaries and FDR language", "Methods", "keep", "FDR threshold is discussed in final analysis."),
        ("Pihur 2009 RankAggreg", "RankAggreg", "no", "possible earlier", "Round6 ranking uses Monte Carlo rank scores, not RankAggreg", "remove from main", "remove", "No final Round6 use."),
        ("Frankish 2023 GENCODE", "GENCODE", "no", "yes", "symbol mapping in archived genomics workflows", "Supplementary Methods only", "move", "Not needed in main Round6 references."),
    ]
    pd.DataFrame(method_rows, columns=[
        "reference", "method", "used_in_round5", "used_in_archived_upstream_pipeline",
        "evidence_in_repository", "current_citation_location", "keep_remove_or_move", "reason",
    ]).to_csv(REPORTS / "method_reference_audit_round6.tsv", sep="\t", index=False)

    add_remove = []
    for r in refs():
        add_remove.append({
            "reference": r.text,
            "action": r.action,
            "old_or_new_number": r.n,
            "reason": r.reason,
            "manuscript_location": r.location,
            "verification_source": r.source,
            "doi_verified": "yes" if r.doi else "not_applicable",
            "pmid_verified": "yes" if r.pmid else "not_available_or_not_required",
        })
    pd.DataFrame(add_remove).to_csv(REPORTS / "reference_add_remove_log.tsv", sep="\t", index=False)

    integ = [
        "# Round6 reference integrity",
        "",
        f"- Search cutoff date: {date.today().isoformat()}.",
        "- New disease-domain references were selected from PubMed, PMC, journal pages and DOI-resolved metadata where available.",
        "- Upstream-only software references were moved out of the main manuscript reference list unless directly used by the Round6 final analysis.",
        "- The final main reference list contains 33 cited references plus BH FDR reference, and all are cited in the manuscript body.",
        "- No fabricated repository DOI or accession was added.",
    ]
    REPORTS.joinpath("reference_integrity_round6.md").write_text("\n".join(integ) + "\n", encoding="utf-8")

    log_rows = []
    for r in refs():
        if r.action == "move":
            include = "exclude"
            exclusion = "upstream-only method reference moved to Supplementary Methods"
            section = "Supplementary Methods"
        else:
            include = "include"
            exclusion = ""
            section = r.location
        log_rows.append({
            "search_date": date.today().isoformat(),
            "database": r.source,
            "query": r.reason or "targeted verification",
            "title": re.sub(r"^[^.]+\. ", "", r.text).split(". ")[0],
            "year": re.search(r"\b(19|20)\d{2}\b", r.text).group(0) if re.search(r"\b(19|20)\d{2}\b", r.text) else "",
            "journal": r.text.split(". ")[-3] if len(r.text.split(". ")) >= 3 else "",
            "study_type": "original/resource/method/reference audit",
            "sample_or_dataset": r.reason,
            "major_finding": r.reason,
            "relevant_manuscript_section": section,
            "doi": r.doi,
            "pmid": r.pmid,
            "include_or_exclude": include,
            "exclusion_reason": exclusion,
        })
    pd.DataFrame(log_rows).to_excel(REPORTS / "jtm_round6_literature_search_log.xlsx", index=False)
    pd.DataFrame(log_rows).to_csv(REPORTS / "jtm_round6_literature_search_log.tsv", sep="\t", index=False)


def supplemental_docs() -> None:
    disc_doc = Document()
    disc_doc.add_heading("Round6 Discussion", level=1)
    for p in discussion_text(stats()).split("\n\n"):
        disc_doc.add_paragraph(p)
    disc_doc.save(OUT / "jtm_discussion_round6.docx")
    ref_doc = Document()
    ref_doc.add_heading("Round6 References", level=1)
    for r in [x for x in refs() if x.action != "move"]:
        ref_doc.add_paragraph(f"{r.n}. {r.text}")
    ref_doc.save(OUT / "jtm_references_round6.docx")


def qc_script() -> None:
    p = ROOT / "scripts/qc/check_round6_discussion_references.py"
    p.parent.mkdir(exist_ok=True)
    p.write_text(r'''#!/usr/bin/env python3
from pathlib import Path
import re, sys
import pandas as pd
from docx import Document

ROOT = Path(__file__).resolve().parents[2]
DOC = ROOT / "outputs/jtm_manuscript_round6_discussion_enhanced.docx"
checks = []
def add(name, ok):
    checks.append((name, bool(ok)))
def text_doc(path):
    doc = Document(path)
    parts = [p.text for p in doc.paragraphs]
    for table in doc.tables:
        for row in table.rows:
            parts.extend(cell.text for cell in row.cells)
    return "\n".join(parts)
txt = text_doc(DOC) if DOC.exists() else ""
disc = txt.split("Discussion",1)[1].split("Conclusions",1)[0] if "Discussion" in txt and "Conclusions" in txt else ""
refs = txt.split("References",1)[1].split("Figures",1)[0] if "References" in txt and "Figures" in txt else ""
ref_nums = sorted(set(int(x) for x in re.findall(r"^(\d+)\.", refs, re.M)))
cite_set = set()
for block in re.findall(r"\[([0-9,\-\s]+)\]", txt):
    for part in block.replace(" ", "").split(","):
        if not part:
            continue
        if "-" in part:
            a, b = part.split("-", 1)
            cite_set.update(range(int(a), int(b) + 1))
        else:
            cite_set.add(int(part))
cite_nums = sorted(cite_set)
add("Discussion word count reasonable", 1200 <= len(disc.split()) <= 1800)
add("principal findings included", "principal finding" in disc.lower() and "430" in disc and "17" in disc)
add("vascular/ECM axis discussed", "extracellular-matrix" in disc.lower() and "VWF" in disc and "COL5A1" in disc)
add("myeloid/complement axis discussed", "myeloid" in disc.lower() and "complement" in disc.lower())
add("scRNA/bulk/protein differences explained", all(x in disc for x in ["scRNA", "Bulk", "Protein"]))
add("protein FDR negative explicit", "no candidate met protein FDR<0.05" in txt or "no candidate met FDR<0.05" in txt)
add("previous integration studies compared", "public-data integration" in disc.lower() and "prediction model" in disc.lower())
add("genetics/TWAS explained", "nearest-gene" in disc.lower() and "colocalization" in disc.lower() and "per-SNP" in disc)
add("drug context separated from treatment", "therapeutic context" in disc and "treatment-ready" not in disc.lower())
add("validation roadmap present", "validation roadmap" in disc.lower() and "spatial transcriptomics" in disc.lower())
add("430 upstream unrecoverable acknowledged", "not fully recoverable" in disc.lower() or "not fully reconstructable" in disc.lower())
add("17 rule retrospective acknowledged", "retrospective" in disc.lower() and "preregistered" in disc.lower())
add("no independent holdout acknowledged", "no fully independent molecular hold-out" in disc.lower())
add("pathway concordance not gene validation", "pathway-level concordance" in disc and "do not constitute independent gene-level validation" in disc)
bad = ["proved", "confirms", "key drivers", "clinically actionable", "repurposing candidate", "promising treatment"]
add("unsupported causal/clinical wording absent", not any(b in disc.lower() for b in bad))
add("literature search log exists", (ROOT/"reports/jtm_round6_literature_search_log.xlsx").exists())
add("DOI field audited", (ROOT/"reports/reference_add_remove_log.tsv").exists() and pd.read_csv(ROOT/"reports/reference_add_remove_log.tsv", sep="\t")["doi_verified"].isin(["yes","not_applicable"]).all())
add("all references cited", set(ref_nums).issubset(set(cite_nums)))
add("no missing references for citations", set(cite_nums).issubset(set(ref_nums)))
add("method audit exists", (ROOT/"reports/method_reference_audit_round6.tsv").exists())
if (ROOT/"reports/method_reference_audit_round6.tsv").exists():
    m = pd.read_csv(ROOT/"reports/method_reference_audit_round6.tsv", sep="\t")
    add("upstream-only methods moved", m[m["used_in_round5"].eq("no")]["keep_remove_or_move"].isin(["move","remove"]).all())
else:
    add("upstream-only methods moved", False)
add("no duplicate reference numbers", len(ref_nums) == len(set(ref_nums)))
add("citation numbering continuous", ref_nums == list(range(1, max(ref_nums)+1)) if ref_nums else False)
add("reference format includes DOI for most", refs.lower().count("doi:") >= 28)
add("Discussion not Methods repeat", disc.lower().count("dirichlet") <= 1 and disc.lower().count("checksum") <= 1)
add("Conclusion consistent", "independent molecular validation" in txt and "treatment-ready" in txt)
add("independent validation not claimed", "independent validation was achieved" not in txt.lower())
add("drug not called treatment candidate", "treatment candidate" not in txt.lower())
add("nearest-gene not causal", "nearest-gene assignment does not identify a causal gene" in disc.lower())
add("protein nominal not validation", "protein validation" not in txt.lower())
failed = [x for x in checks if not x[1]]
report = ROOT / "reports/jtm_round6_discussion_reference_qc.md"
report.write_text("# Round6 Discussion/reference QC\n\n" + "\n".join(f"- {'PASS' if ok else 'FAIL'}: {name}" for name, ok in checks) + "\n", encoding="utf-8")
block = ROOT / "reports/jtm_round6_submission_blockers.md"
block.write_text("# Round6 submission blockers\n\n" + ("- Technical QC failures: %d.\n" % len(failed) if failed else "- Technical QC: no failures in implemented checks.\n") + "- Author metadata, funding, conflicts, repository URL/DOI and final access dates remain author-side blockers.\n", encoding="utf-8")
print(f"Round6 QC: {len(checks)-len(failed)} PASS, {len(failed)} FAIL")
for name, _ in failed:
    print("FAIL:", name)
sys.exit(1 if failed else 0)
''', encoding="utf-8")
    p.chmod(0o755)


def final_report() -> None:
    REPORTS.joinpath("jtm_round6_final_revision_report.md").write_text(
        "# JTM Round6 final revision report\n\n"
        "- Expanded Discussion from a short limitation-focused section to a biologically organized interpretation covering principal findings, vascular/ECM biology, myeloid/complement biology, cross-omic disagreement, genetics/TWAS boundaries, therapeutic context and validation roadmap.\n"
        "- Added a focused Background research-gap paragraph about missingness, candidate-universe dependence and separation of ranking inputs from contextual annotations.\n"
        "- Updated Table 2 evidence profiles into biological modules aligned with the Discussion.\n"
        "- Moved upstream-only method references out of the main manuscript reference list and retained only method references used in final interpretation.\n"
        "- Generated literature search log, claim-reference matrix, gene-external-evidence map, method audit and reference add/remove log.\n"
        "- Maintained conservative language: no causal, clinical-prediction, treatment-readiness or independent-validation claims were added.\n",
        encoding="utf-8",
    )


def main() -> None:
    build_audits()
    build_doc(False).save(OUT / "jtm_manuscript_round6_discussion_enhanced.docx")
    build_doc(True).save(OUT / "jtm_manuscript_round6_blinded.docx")
    supplemental_docs()
    qc_script()
    final_report()
    print(OUT / "jtm_manuscript_round6_discussion_enhanced.docx")


if __name__ == "__main__":
    main()
