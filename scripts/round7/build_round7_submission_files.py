#!/usr/bin/env python3
from __future__ import annotations

from pathlib import Path
from docx import Document
from docx.shared import Pt

ROOT = Path(__file__).resolve().parents[2]
OUT = ROOT / "outputs"
REPORTS = ROOT / "reports"
OUT.mkdir(exist_ok=True)
REPORTS.mkdir(exist_ok=True)

TITLE = "Retrospective multi-source evidence audit and candidate prioritization in proliferative diabetic retinopathy"
AUTHORS_LINE = "Yi Zha1, Da Lin1, Ying Chen2, Yue Liu2, Yu Zhang1,*"
AFFILIATIONS = [
    "1 The Second Affiliated Hospital of Wenzhou Medical University, Wenzhou, Zhejiang, China.",
    "2 Wenzhou Medical University, Wenzhou, Zhejiang, China.",
]
CORRESPONDENCE = (
    "*Correspondence: Yu Zhang, The Second Affiliated Hospital of Wenzhou Medical University, "
    "Wenzhou, Zhejiang, China. Email: zhangyu1@wzhealth.com. "
    "ORCID: 0000-0001-8579-3692."
)
GITHUB = "https://github.com/seefreewind/pdr-public-multiomics-prioritization"
ZENODO = "https://doi.org/10.5281/zenodo.21404212"


DECLARATIONS = {
    "Acknowledgements": (
        "The authors thank the investigators, data contributors and participants associated with the public datasets "
        "used in this study."
    ),
    "Authors' contributions": (
        "Yi Zha, Da Lin, Ying Chen, Yue Liu and Yu Zhang contributed to study conception, data curation, analysis, "
        "interpretation and manuscript revision. Yu Zhang supervised the project. All authors read and approved the "
        "final manuscript. [AUTHOR CONFIRMATION REQUIRED: please verify individual contribution details before submission.]"
    ),
    "Funding": (
        "[AUTHOR CONFIRMATION REQUIRED: insert grant numbers and funder names if applicable. If no specific funding "
        "supported this work, replace this sentence with: This research received no specific grant from any funding "
        "agency in the public, commercial or not-for-profit sectors.]"
    ),
    "Availability of data and materials": (
        f"The analysis code and processed, non-restricted manuscript-support files are available at GitHub: {GITHUB}. "
        f"The version corresponding to this manuscript has been archived at Zenodo: {ZENODO}. Raw public datasets "
        "should be obtained from the original repositories using the accessions reported in the manuscript."
    ),
    "Ethics approval and consent to participate": (
        "This study was based on publicly available and previously generated datasets. No new human participants were "
        "recruited, and no new biospecimens were collected for this secondary analysis."
    ),
    "Consent for publication": "Not applicable.",
    "Competing interests": (
        "[AUTHOR CONFIRMATION REQUIRED: if accurate, replace this sentence with: The authors declare that they have no competing interests.]"
    ),
}


def set_para_text(paragraph, text: str) -> None:
    paragraph.clear()
    run = paragraph.add_run(text)
    run.font.name = "Arial"
    run.font.size = Pt(10)


def finalize_main_doc() -> Path:
    src = OUT / "jtm_manuscript_round6_discussion_enhanced.docx"
    doc = Document(src)
    # Replace author placeholder under title.
    for i, p in enumerate(doc.paragraphs):
        if p.text.strip() == "[AUTHOR CONFIRMATION REQUIRED]":
            set_para_text(p, AUTHORS_LINE + "\n" + "\n".join(AFFILIATIONS) + "\n" + CORRESPONDENCE)
            break
    # Replace declaration placeholder paragraphs by heading.
    for i, p in enumerate(doc.paragraphs):
        heading = p.text.strip()
        if heading in DECLARATIONS:
            j = i + 1
            while j < len(doc.paragraphs) and doc.paragraphs[j].text.strip() == "":
                j += 1
            if j < len(doc.paragraphs):
                set_para_text(doc.paragraphs[j], DECLARATIONS[heading])
    out = OUT / "jtm_manuscript_round7_submission_ready_pending_author_confirmation.docx"
    doc.save(out)
    return out


def title_page() -> Path:
    doc = Document()
    doc.add_heading("Title Page", level=1)
    doc.add_paragraph(TITLE)
    doc.add_heading("Authors", level=2)
    doc.add_paragraph(AUTHORS_LINE)
    doc.add_heading("Affiliations", level=2)
    for a in AFFILIATIONS:
        doc.add_paragraph(a)
    doc.add_heading("Corresponding Author", level=2)
    doc.add_paragraph(CORRESPONDENCE)
    doc.add_heading("Repository and Archive", level=2)
    doc.add_paragraph(f"GitHub: {GITHUB}")
    doc.add_paragraph(f"Zenodo: {ZENODO}")
    doc.add_heading("Short Title", level=2)
    doc.add_paragraph("PDR multi-source evidence audit")
    doc.add_heading("Manuscript Type", level=2)
    doc.add_paragraph("Research article")
    out = OUT / "jtm_title_page_round7.docx"
    doc.save(out)
    return out


def declarations_doc() -> Path:
    doc = Document()
    doc.add_heading("Declarations", level=1)
    for h, text in DECLARATIONS.items():
        doc.add_heading(h, level=2)
        doc.add_paragraph(text)
    out = OUT / "jtm_declarations_round7_pending_author_confirmation.docx"
    doc.save(out)
    return out


def cover_letter() -> Path:
    doc = Document()
    doc.add_paragraph("Dear Editors,")
    doc.add_paragraph(
        "We are pleased to submit our manuscript entitled "
        f"\"{TITLE}\" for consideration as a Research Article in Journal of Translational Medicine."
    )
    doc.add_paragraph(
        "This study presents a retrospective public-data evidence audit and secondary prioritization of a frozen "
        "430-gene candidate universe in proliferative diabetic retinopathy. The work reconstructs the fixed candidate "
        "universe, reproduces a focused 17-gene prioritized set from archived rules, distinguishes ranking inputs from "
        "contextual annotations, and evaluates sensitivity to evidence-layer composition."
    )
    doc.add_paragraph(
        "The manuscript emphasizes conservative interpretation. It does not claim same-patient multimodal fusion, "
        "independent molecular validation, causal genes, clinical biomarkers or treatment-ready interventions. Instead, "
        "it provides an auditable framework for interpreting public multi-source evidence and identifying candidates "
        "for future tissue, cellular, genetic and imaging-linked validation."
    )
    doc.add_paragraph(
        f"The analysis code and processed, non-restricted manuscript-support files are available at GitHub ({GITHUB}), "
        f"and the submitted version has been archived at Zenodo ({ZENODO})."
    )
    doc.add_paragraph(
        "All authors should confirm the final funding, competing-interest and contribution statements before submission."
    )
    doc.add_paragraph("Sincerely,")
    doc.add_paragraph("Yu Zhang\nOn behalf of all authors")
    out = OUT / "jtm_cover_letter_round7.docx"
    doc.save(out)
    return out


def confirmation_form() -> Path:
    doc = Document()
    doc.add_heading("Round7 Author Confirmation Checklist", level=1)
    checks = [
        "Confirm author order: Yi Zha, Da Lin, Ying Chen, Yue Liu, Yu Zhang.",
        "Confirm affiliations for all authors.",
        "Confirm corresponding author email and ORCID for Yu Zhang.",
        "Confirm individual author contributions.",
        "Confirm funding statement and grant numbers, or confirm no specific funding.",
        "Confirm competing-interest statement.",
        "Confirm that GitHub and Zenodo records may be cited in the manuscript.",
        "Confirm whether GitHub should remain private or be made public before submission.",
        "Confirm Zenodo access status and license.",
    ]
    for c in checks:
        doc.add_paragraph("[  ] " + c)
    out = OUT / "jtm_author_confirmation_checklist_round7.docx"
    doc.save(out)
    return out


def report(paths: list[Path]) -> Path:
    remaining = [
        "Funding statement requires author confirmation.",
        "Competing-interest statement requires author confirmation.",
        "Detailed author-contribution statement requires author confirmation.",
        "GitHub repository visibility and Zenodo license/access status should be confirmed before formal submission.",
    ]
    lines = [
        "# Round7 submission finalization report",
        "",
        "Completed:",
        "- Inserted author line, affiliations, corresponding author email and ORCID into the manuscript.",
        "- Inserted GitHub and Zenodo DOI into data availability wording.",
        "- Added public-data ethics/consent wording.",
        "- Generated title page, cover letter, declarations document and author confirmation checklist.",
        "- Preserved conservative claims and did not add causal, clinical-validation or treatment-readiness wording.",
        "",
        "Generated files:",
    ]
    lines.extend(f"- {p.relative_to(ROOT)}" for p in paths)
    lines += ["", "Remaining author-side confirmations:"]
    lines.extend(f"- {x}" for x in remaining)
    out = REPORTS / "jtm_round7_submission_finalization_report.md"
    out.write_text("\n".join(lines) + "\n", encoding="utf-8")
    return out


def main() -> None:
    paths = [finalize_main_doc(), title_page(), declarations_doc(), cover_letter(), confirmation_form()]
    paths.append(report(paths))
    print("\n".join(str(p) for p in paths))


if __name__ == "__main__":
    main()
